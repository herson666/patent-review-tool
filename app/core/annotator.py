"""将 FormIssue 列表写入 Word 文档批注"""
import os
from typing import List
from docx import Document
from docx.shared import Pt, RGBColor
from app.models import FormIssue, Severity, IssueSource


AUTHOR_NAME = "专利形式审核"


def annotate_document(doc_path: str, output_path: str, issues: List[FormIssue],
                      all_paragraphs=None) -> int:
    """
    将 FormIssue 列表写入 Word 文档批注。
    返回批注数量。
    """
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"源文件不存在: {doc_path}")

    doc = Document(doc_path)
    _clear_existing_comments(doc)

    comment_count = 0
    for issue in issues:
        para_idx = issue.range.paragraph_index
        if para_idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[para_idx]
        # 构建批注文本
        source_label = _get_source_label(issue.source)
        lines = []
        if issue.rule_id:
            lines.append(f"[{issue.rule_id}]")
        lines.append(f"[{source_label}]")
        lines.append(issue.message)
        if issue.law_reference:
            lines.append(f"法规依据: {issue.law_reference}")
        comment_text = "\n".join(lines)

        try:
            para.add_comment(comment_text, author=AUTHOR_NAME, initials="ZLS")
            comment_count += 1
        except (AttributeError, TypeError) as e:
            # python-docx 版本不支持 add_comment，fallback：添加带颜色的内联文本
            _add_fallback_annotation(para, comment_text, issue)

    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return comment_count


def _clear_existing_comments(doc: Document):
    """清除文档中所有已有批注（避免重复）"""
    try:
        comments_part = doc.part.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
        )
        if comments_part:
            # 重置批注内容
            from docx.oxml import parse_xml
            empty = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'''
            comments_part._blob = empty.encode('utf-8')
    except Exception:
        pass


def _get_source_label(source: IssueSource) -> str:
    if source == IssueSource.RULE_ENGINE:
        return "规则"
    elif source == IssueSource.MODEL_FORM:
        return "模型-形式"
    elif source == IssueSource.MODEL_FLUENCY:
        return "模型-通顺"
    return "未知"


def _get_color(severity: Severity, source: IssueSource) -> RGBColor:
    """根据严重程度和来源返回颜色"""
    if source == IssueSource.MODEL_FLUENCY:
        return RGBColor(0xFF, 0xA5, 0x00)  # 橙色
    if severity == Severity.ERROR:
        return RGBColor(0xFF, 0x00, 0x00)  # 红色
    return RGBColor(0xFF, 0x8C, 0x00)  # 深橙色


def _add_fallback_annotation(para, text: str, issue: FormIssue):
    """fallback: 在段落末尾添加带颜色的注释文字"""
    color = _get_color(issue.severity, issue.source)
    run = para.add_run(f" 【{text}】")
    run.font.color.rgb = color
    run.font.size = Pt(9)
