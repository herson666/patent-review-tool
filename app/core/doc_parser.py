"""专利申请文件 .docx 解析与节识别"""
import re
from typing import List, Optional
from docx import Document
from docx.oxml.ns import qn
from app.models import DocumentModel, Section, SectionType


SECTION_KEYWORDS = {
    SectionType.ABSTRACT: ["摘要"],
    SectionType.CLAIMS: ["权利要求"],
    SectionType.FIGURES: ["说明书附图", "附图"],
    SectionType.DESCRIPTION: ["说明书"],
}


def _identify_section_type(header_text: str, first_para: str) -> SectionType:
    """根据页眉和首页首段识别节类型

    优先级：页眉 > 首段文本
    """
    # 1) 页眉优先（最可靠）
    header = (header_text or "").strip()
    if header:
        if "说明书摘要" in header or header == "摘要":
            return SectionType.ABSTRACT
        if "摘要附图" in header:
            # 摘要附图按附图处理（保持与说明书附图一致类型）
            return SectionType.FIGURES
        if "权利要求书" in header or "权利要求" in header:
            return SectionType.CLAIMS
        if "说明书附图" in header or header == "附图":
            return SectionType.FIGURES
        if "说明书" in header:
            return SectionType.DESCRIPTION

    # 2) 首段文本兜底
    text = (header_text or "") + "\n" + (first_para or "")
    # 摘要：含"摘要"但不含"附图"
    if "摘要" in text and "附图" not in text and "权利要求" not in text:
        return SectionType.ABSTRACT
    if "权利要求" in text:
        return SectionType.CLAIMS
    if "附图" in text or re.match(r"^图\s*\d+", (first_para or "").strip()):
        return SectionType.FIGURES
    if "说明书" in text or "技术领域" in text or "发明内容" in text or "具体实施方式" in text:
        return SectionType.DESCRIPTION
    return SectionType.UNKNOWN


def _get_section_header_text(doc: Document, section_idx: int) -> str:
    """获取指定 docx section 的页眉文本"""
    if section_idx < 0 or section_idx >= len(doc.sections):
        return ""
    sec = doc.sections[section_idx]
    parts = []
    for p in sec.header.paragraphs:
        if p.text:
            parts.append(p.text)
    return "\n".join(parts)


def _split_by_section_breaks(doc: Document) -> List[List]:
    """按 docx 的真实分节符（sectPr）切分段落列表

    返回：每节对应的 python-docx Paragraph 对象列表
    """
    body = doc.element.body
    section_groups: List[List] = []
    current: List = []

    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, doc.part)
            current.append(para)
            # 分节符位于段落 pPr 内的 sectPr
            sectPr = child.find(qn("w:pPr") + "/" + qn("w:sectPr"))
            if sectPr is not None:
                section_groups.append(current)
                current = []
        elif tag == qn("w:tbl"):
            # 跳过表格行（CLAIMS 表格行已在 e2e 场景下处理；这里按段落处理时不展开）
            continue
        # 忽略其他元素

    if current:
        section_groups.append(current)

    return section_groups


def parse_docx(file_path: str) -> DocumentModel:
    """
    解析 .docx 文件，自动识别五书节（摘要、权利要求书、说明书、说明书附图）。

    节识别策略：
    1. 优先按 docx 真实分节符（sectPr）切分
    2. 提取每节页眉作为类型判定的首要依据
    3. 结合页眉与首段非空文本判断 SectionType
    4. 若分节符缺失或页眉为空，回退到按段落内容扫描的旧逻辑
    """
    doc = Document(file_path)
    all_paragraphs = list(doc.paragraphs)
    all_texts = [p.text for p in all_paragraphs]

    # 提取发明名称（文档首段）
    patent_name = all_texts[0].strip() if all_texts else ""

    # 1) 按分节符分节
    section_groups = _split_by_section_breaks(doc)

    sections: List[Section] = []
    if len(section_groups) > 1:
        # 多个 docx section：使用页眉 + 首段判定
        cursor = 0
        for i, paras in enumerate(section_groups):
            header_text = _get_section_header_text(doc, i)
            first_para = next((p.text for p in paras if (p.text or "").strip()), "")
            stype = _identify_section_type(header_text, first_para)
            if stype == SectionType.UNKNOWN:
                # 兜底：用首段或页眉
                stype = _identify_section_type("", first_para)
            texts = [p.text for p in paras]
            if stype != SectionType.UNKNOWN:
                sections.append(Section(
                    section_type=stype,
                    title=texts[0][:50] if texts else "",
                    paragraphs=texts,
                    raw_paragraphs=paras,
                    start_para_index=cursor
                ))
            cursor += len(paras)
    else:
        # 单节文件：使用旧的关键词扫描
        sections = _legacy_scan_sections(all_paragraphs, all_texts)

    return DocumentModel(
        file_path=file_path,
        sections=sections,
        patent_name=patent_name,
        all_paragraphs=all_paragraphs
    )


def _legacy_scan_sections(all_paragraphs, all_texts) -> List[Section]:
    """旧的关键词扫描分节逻辑（无分节符时回退）"""
    sections: List[Section] = []
    current_type = SectionType.UNKNOWN
    current_paras: List[str] = []
    current_raw: List = []
    current_start = 0

    for i, (text, raw) in enumerate(zip(all_texts, all_paragraphs)):
        identified = _identify_section_type("", text)
        if identified != SectionType.UNKNOWN and identified != current_type:
            if current_paras and current_type != SectionType.UNKNOWN:
                sections.append(Section(
                    section_type=current_type,
                    title=current_paras[0][:50] if current_paras else "",
                    paragraphs=current_paras,
                    raw_paragraphs=current_raw,
                    start_para_index=current_start
                ))
            current_type = identified
            current_paras = [text]
            current_raw = [raw]
            current_start = i
        else:
            current_paras.append(text)
            current_raw.append(raw)

    if current_paras and current_type != SectionType.UNKNOWN:
        sections.append(Section(
            section_type=current_type,
            title=current_paras[0][:50] if current_paras else "",
            paragraphs=current_paras,
            raw_paragraphs=current_raw,
            start_para_index=current_start
        ))

    return sections


def get_section_by_type(model: DocumentModel, section_type: SectionType) -> Optional[Section]:
    """从 DocumentModel 中获取指定类型的节"""
    for section in model.sections:
        if section.section_type == section_type:
            return section
    return None


def get_all_figure_refs(model: DocumentModel) -> set:
    """获取文档中所有附图标记（用于一致性检查）"""
    refs = set()
    for section in model.sections:
        for para in section.paragraphs:
            refs.update(re.findall(r"图(\d+)", para))
    return refs
