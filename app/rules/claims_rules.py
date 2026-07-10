"""权利要求书形式审核规则（11条）"""
import re
from typing import List
from app.models import Section, FormIssue, TextRange, Severity, SectionType, IssueSource
from app.rules.base import BaseRule


class Claims001Rule(BaseRule):
    """R1: 多项权利要求应当用阿拉伯数字顺序编号"""
    section_type = SectionType.CLAIMS

    @property
    def description(self) -> str:
        return "多项权利要求应当用阿拉伯数字顺序编号"

    @property
    def law_reference(self) -> str:
        return "实施细则第22条"

    def check(self, section: Section) -> List[FormIssue]:
        issues = []
        # 识别可能的权利要求段落
        claim_indicators = ["权利要求", "如权利要求", "根据权利要求"]
        numbered_pattern = re.compile(r"^\s*(\d+)[\.、]")

        for i, para in enumerate(section.paragraphs):
            text = para.strip()
            if not text:
                continue
            # 检查是否含权利要求指示词
            has_indicator = any(ind in text for ind in claim_indicators)
            if not has_indicator:
                continue
            # 检查是否已编号
            if not numbered_pattern.match(text):
                issues.append(FormIssue(
                    section=section.section_type,
                    range=TextRange(paragraph_index=i),
                    severity=Severity.ERROR,
                    rule_id="CLAIMS-001",
                    message="权利要求应当用阿拉伯数字顺序编号",
                    law_reference=self.law_reference,
                    source=IssueSource.RULE_ENGINE
                ))
        return issues


class Claims002Rule(BaseRule):
    """R2: 从属权利要求应当尽量紧靠其所引用的权利要求"""
    section_type = SectionType.CLAIMS

    @property
    def description(self) -> str:
        return "从属权利要求应当尽量紧靠其所引用的权利要求"

    @property
    def law_reference(self) -> str:
        return "实施细则第22条"

    def check(self, section: Section) -> List[FormIssue]:
        # 简化实现：检测"根据权利要求X所述的"出现在独立权利要求之后的段落
        issues = []
        ref_pattern = re.compile(r"根据权利要求(\d+)所述")
        numbered_pattern = re.compile(r"^\s*(\d+)[\.、]")
        claim_positions = {}  # 编号 -> 段落索引

        for i, para in enumerate(section.paragraphs):
            m = numbered_pattern.match(para.strip())
            if m:
                claim_positions[int(m.group(1))] = i

        for i, para in enumerate(section.paragraphs):
            text = para.strip()
            refs = ref_pattern.findall(text)
            for ref_str in refs:
                ref_num = int(ref_str)
                if ref_num in claim_positions:
                    gap = i - claim_positions[ref_num]
                    if gap > 12:  # 间隔超过 12 段视为距离过大
                        issues.append(FormIssue(
                            section=section.section_type,
                            range=TextRange(paragraph_index=i),
                            severity=Severity.WARNING,
                            rule_id="CLAIMS-002",
                            message=f"从属权利要求引用的权利要求{ref_num}在{claim_positions[ref_num]+1}段，距此段落{gap}段，间隔过大",
                            law_reference=self.law_reference,
                            source=IssueSource.RULE_ENGINE
                        ))
        return issues


class Claims003Rule(BaseRule):
    """R3: 每一项权利要求只允许在其结尾使用句号"""
    section_type = SectionType.CLAIMS

    @property
    def description(self) -> str:
        return "每一项权利要求只允许在其结尾使用句号"

    @property
    def law_reference(self) -> str:
        return "实施细则第22条"

    def check(self, section: Section) -> List[FormIssue]:
        issues = []
        numbered_pattern = re.compile(r"^\s*\d+[\.、]")
        for i, para in enumerate(section.paragraphs):
            text = para.strip()
            if not text:
                continue
            # 识别权利要求项：已编号 或 含"权利要求"指示词
            is_claim = bool(numbered_pattern.match(text)) or "权利要求" in text
            if not is_claim:
                continue
            # 句号出现在非结尾位置
            if "。" in text[:-1]:
                issues.append(FormIssue(
                    section=section.section_type,
                    range=TextRange(paragraph_index=i),
                    severity=Severity.WARNING,
                    rule_id="CLAIMS-003",
                    message="权利要求中除结尾外不得使用句号",
                    law_reference=self.law_reference,
                    source=IssueSource.RULE_ENGINE
                ))
        return issues


class Claims004Rule(BaseRule):
    """R4: 权利要求中使用的科技术语应当与说明书中一致"""
    section_type = SectionType.CLAIMS

    @property
    def description(self) -> str:
        return "权利要求中使用的科技术语应当与说明书中一致"

    @property
    def law_reference(self) -> str:
        return "实施细则第22条"

    def check(self, section: Section) -> List[FormIssue]:
        # 此规则需要结合说明书，当前简化实现
        return []


class Claims005Rule(BaseRule):
    """R5: 权利要求中不得有插图"""
    section_type = SectionType.CLAIMS

    @property
    def description(self) -> str:
        return "权利要求中可以有化学式或数学式，但不得有插图"

    @property
    def law_reference(self) -> str:
        return "实施细则第22条"

    def check(self, section: Section) -> List[FormIssue]:
        return []


class Claims006Rule(BaseRule):
    """R6: 除非绝对必要，不得使用引用语"""
    section_type = SectionType.CLAIMS

    @property
    def description(self) -> str:
        return "除非绝对必要，不得使用如说明书…所述的引用语"

    @property
    def law_reference(self) -> str:
        return "实施细则第22条"

    def check(self, section: Section) -> List[FormIssue]:
        issues = []
        pattern = re.compile(
            r"如.{0,15}(?:说明书|权利要求|附图).{0,30}所述|"
            r"如.{0,10}图.{0,10}所示"
        )
        for i, para in enumerate(section.paragraphs):
            text = para.strip()
            matches = pattern.findall(text)
            if matches:
                issues.append(FormIssue(
                    section=section.section_type,
                    range=TextRange(paragraph_index=i),
                    severity=Severity.WARNING,
                    rule_id="CLAIMS-006",
                    message=f"使用了引用语「{matches[0]}」，除非绝对必要不得使用",
                    law_reference=self.law_reference,
                    source=IssueSource.RULE_ENGINE
                ))
        return issues


class Claims007Rule(BaseRule):
    """R7: 权利要求中通常不允许使用表格"""
    section_type = SectionType.CLAIMS

    @property
    def description(self) -> str:
        return "权利要求中通常不允许使用表格"

    @property
    def law_reference(self) -> str:
        return "实施细则第22条"

    def check(self, section: Section) -> List[FormIssue]:
        issues = []
        for i, raw in enumerate(section.raw_paragraphs):
            from docx.table import Table
            if isinstance(raw, Table):
                issues.append(FormIssue(
                    section=section.section_type,
                    range=TextRange(paragraph_index=i),
                    severity=Severity.WARNING,
                    rule_id="CLAIMS-007",
                    message="权利要求中通常不允许使用表格",
                    law_reference=self.law_reference,
                    source=IssueSource.RULE_ENGINE
                ))
        return issues


class Claims008Rule(BaseRule):
    """R8: 附图标记应当放在括号内"""
    section_type = SectionType.CLAIMS

    @property
    def description(self) -> str:
        return "权利要求中的附图标记应当放在括号内"

    @property
    def law_reference(self) -> str:
        return "实施细则第22条"

    def check(self, section: Section) -> List[FormIssue]:
        issues = []
        # 检测 图1 图2 等出现在非括号中（同时支持中英文括号）
        fig_pattern = re.compile(r"(?<![(\（])图\d+(?![)\）])")
        for i, para in enumerate(section.paragraphs):
            matches = fig_pattern.findall(para)
            if matches:
                issues.append(FormIssue(
                    section=section.section_type,
                    range=TextRange(paragraph_index=i),
                    severity=Severity.ERROR,
                    rule_id="CLAIMS-008",
                    message=f"附图标记「{matches[0]}」应加括号",
                    law_reference=self.law_reference,
                    source=IssueSource.RULE_ENGINE
                ))
        return issues


class Claims009Rule(BaseRule):
    """R9: 附图标记不得解释为对权利要求保护范围的限制"""
    section_type = SectionType.CLAIMS

    @property
    def description(self) -> str:
        return "附图标记不得解释为对权利要求保护范围的限制"

    @property
    def law_reference(self) -> str:
        return "实施细则第22条"

    def check(self, section: Section) -> List[FormIssue]:
        return []  # 声明性规则，由审查员判断


class Claims010Rule(BaseRule):
    """R10: 除附图标记必要情形外，应当尽量避免使用括号"""
    section_type = SectionType.CLAIMS

    @property
    def description(self) -> str:
        return "除附图标记或其它必要情形必须使用括号外，权利要求中应当尽量避免使用括号"

    @property
    def law_reference(self) -> str:
        return "实施细则第22条"

    def check(self, section: Section) -> List[FormIssue]:
        issues = []
        # 同时支持中英文括号
        bracket_pattern = re.compile(r"（[^）]+）|\([^)]+\)")
        for i, para in enumerate(section.paragraphs):
            brackets = bracket_pattern.findall(para)
            non_fig = [b for b in brackets if not re.search(r"图\d+", b)]
            if non_fig:
                issues.append(FormIssue(
                    section=section.section_type,
                    range=TextRange(paragraph_index=i),
                    severity=Severity.WARNING,
                    rule_id="CLAIMS-010",
                    message="权利要求中除附图标记外应尽量避免使用括号",
                    law_reference=self.law_reference,
                    source=IssueSource.RULE_ENGINE
                ))
        return issues


class Claims011Rule(BaseRule):
    """R11: 从属权利要求只能引用在前的权利要求"""
    section_type = SectionType.CLAIMS

    @property
    def description(self) -> str:
        return "从属权利要求只能引用在前的权利要求"

    @property
    def law_reference(self) -> str:
        return "实施细则第25条"

    def check(self, section: Section) -> List[FormIssue]:
        issues = []
        numbered_pattern = re.compile(r"^\s*(\d+)[\.、]")
        claim_nums = {}  # 编号 -> 段落索引

        for i, para in enumerate(section.paragraphs):
            m = numbered_pattern.match(para.strip())
            if m:
                claim_nums[int(m.group(1))] = i

        ref_pattern = re.compile(r"权利要求(\d+)")
        for i, para in enumerate(section.paragraphs):
            refs = ref_pattern.findall(para)
            for ref_str in refs:
                ref_num = int(ref_str)
                if ref_num not in claim_nums:
                    issues.append(FormIssue(
                        section=section.section_type,
                        range=TextRange(paragraph_index=i),
                        severity=Severity.ERROR,
                        rule_id="CLAIMS-011",
                        message=f"引用了不存在的权利要求{ref_num}",
                        law_reference=self.law_reference,
                        source=IssueSource.RULE_ENGINE
                    ))
        return issues
